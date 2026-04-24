#!/usr/bin/env python3
"""
Generate the Step 7 review cover memo as DOCX, with Markdown fallback.

Usage:
    python3 generate-cover-memo.py <review_manifest_json> <issue_registry_json>
        <review_scorecard_json> <verification_audit_json> <output_docx>
"""

import json
import os
import sys
from datetime import datetime


def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def severity_title(value):
    raw = str(value or "Minor").strip().lower()
    if raw.startswith("crit"):
        return "Critical"
    if raw.startswith("maj"):
        return "Major"
    if raw.startswith("sug"):
        return "Suggestion"
    return "Minor"


def dedupe_recommendations(issues, limit=5):
    seen = set()
    recommendations = []
    for issue in issues:
        text = str(issue.get("recommendation") or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        recommendations.append(text)
        if len(recommendations) >= limit:
            break
    return recommendations


def build_sections(manifest, issue_registry, review_scorecard, verification_audit):
    language = manifest.get("document", {}).get("language", "ko")
    is_korean = str(language).lower().startswith("ko")
    document = manifest.get("document", {})
    issues = issue_registry.get("issues", [])
    release = review_scorecard.get("release_recommendation", "Pass with Warnings")
    grade = review_scorecard.get("overall_grade", "C")
    reviewer = "시니어 리뷰 스페셜리스트" if is_korean else "Senior Review Specialist"
    today = datetime.now().strftime("%Y-%m-%d")

    critical = [issue for issue in issues if severity_title(issue.get("severity")) == "Critical"]
    major = [issue for issue in issues if severity_title(issue.get("severity")) == "Major"]
    minor = [issue for issue in issues if severity_title(issue.get("severity")) in {"Minor", "Suggestion"}]
    recurring = [issue for issue in issues if issue.get("recurring_pattern")]
    citation_count = verification_audit.get("total_citations", len(verification_audit.get("citations", [])))

    if is_korean:
        if not issues:
            overall_assessment = (
                f"○ 이번 검토에서 등록된 이슈는 없습니다. 인용 검증 대상 {citation_count}건 기준으로도 "
                "별도 수정 권고 없이 현재 scorecard의 릴리스 권고를 따를 수 있습니다."
            )
        elif critical:
            overall_assessment = (
                f"○ 총 {len(issues)}건의 이슈 중 Critical {len(critical)}건이 확인되었습니다. "
                "현재 버전은 해당 항목 수정 전 외부 공유를 보류하는 것이 안전합니다."
            )
        elif major:
            overall_assessment = (
                f"○ 총 {len(issues)}건의 이슈 중 Major {len(major)}건이 확인되었습니다. "
                "문서의 기본 구조는 유지 가능하나, 표시된 주요 항목은 릴리스 전 정리하는 것이 좋습니다."
            )
        else:
            overall_assessment = (
                f"○ 총 {len(issues)}건의 Minor/Suggestion 항목이 확인되었습니다. "
                "현재 scorecard상 릴리스를 막는 중대 이슈는 확인되지 않았습니다."
            )
        style_text = "스타일 비교 데이터가 없으면 생략 가능하나, 이번 메모에서는 별도 스타일 프로파일 입력이 없어 비교를 수행하지 않았습니다."
        next_steps_intro = "우선순위는 다음과 같습니다."
        no_next_steps = "현재 issue registry 기준 추가 수정 권고는 없습니다."
        no_critical = "Critical finding 없음"
        no_major = "Major finding 없음"
        no_recurring = "반복 패턴 없음"
    else:
        if not issues:
            overall_assessment = (
                f"No findings were registered in this review. With {citation_count} citations reviewed, "
                "the document can follow the scorecard release recommendation without additional issue-driven revisions."
            )
        elif critical:
            overall_assessment = (
                f"{len(issues)} findings were registered, including {len(critical)} Critical finding(s). "
                "Release should be held until those items are corrected."
            )
        elif major:
            overall_assessment = (
                f"{len(issues)} findings were registered, including {len(major)} Major finding(s). "
                "The document should be revised before external release."
            )
        else:
            overall_assessment = (
                f"{len(issues)} Minor/Suggestion item(s) were registered. "
                "No release-blocking issue appears in the issue registry."
            )
        style_text = "Style fingerprint comparison was skipped because no style profile or sufficient comparison samples were provided."
        next_steps_intro = "Recommended priority order:"
        no_next_steps = "No additional issue-driven revisions are recommended."
        no_critical = "No Critical findings."
        no_major = "No Major findings."
        no_recurring = "No recurring pattern matches."

    def summarize_issue(issue):
        location = issue.get("location", {})
        location_text = location.get("section") or (
            f"para {location.get('paragraph_index')}" if "paragraph_index" in location else "general"
        )
        return f"{location_text}: {issue.get('description')} / {issue.get('recommendation')}"

    recommendations = dedupe_recommendations(critical + major + minor)

    sections = {
        "language": language,
        "title": "문서 검토 결과 보고" if is_korean else "Document Review Report",
        "identification": [
            ("검토 대상" if is_korean else "Document", document.get("title", document.get("filename", ""))),
            ("작성자" if is_korean else "Author", document.get("author", "")),
            ("검토일" if is_korean else "Review Date", today),
            ("검토자" if is_korean else "Reviewer", reviewer),
        ],
        "release_heading": "릴리스 권고" if is_korean else "Release Recommendation",
        "release": release,
        "release_rationale": review_scorecard.get("release_rationale", ""),
        "overall_heading": "전체 평가" if is_korean else "Overall Assessment",
        "overall_assessment": overall_assessment,
        "grade_label": "종합 등급" if is_korean else "Overall Grade",
        "grade": grade,
        "scorecard_heading": "리뷰 스코어카드" if is_korean else "Review Scorecard",
        "critical_heading": "Critical 소견 (반드시 수정)" if is_korean else "Critical Findings (Must Fix)",
        "critical_items": [summarize_issue(issue) for issue in critical] or [no_critical],
        "major_heading": "Major 소견 (수정 권고)" if is_korean else "Major Findings (Should Fix)",
        "major_items": [summarize_issue(issue) for issue in major] or [no_major],
        "minor_heading": "Minor 소견 및 제안" if is_korean else "Minor Findings & Suggestions",
        "minor_items": [summarize_issue(issue) for issue in minor[:5]] or ["없음" if is_korean else "None."],
        "recurring_heading": "반복 패턴 알림" if is_korean else "Recurring Pattern Alerts",
        "recurring_items": [
            f"{issue.get('recurring_pattern')}: {issue.get('description')}" for issue in recurring
        ] or [no_recurring],
        "style_heading": "스타일 비교" if is_korean else "Style Fingerprint Comparison",
        "style_text": style_text,
        "next_heading": "권장 다음 단계" if is_korean else "Recommended Next Steps",
        "next_items": [next_steps_intro, *recommendations] if recommendations else [no_next_steps],
    }
    return sections


def write_markdown_fallback(output_docx, sections, review_scorecard):
    fallback = os.path.splitext(output_docx)[0] + ".md"
    lines = [
        f"# {sections['title']}",
        "",
        f"**{sections['release_heading']}**",
        "",
        f"**{sections['release']}**",
        "",
        sections["release_rationale"],
        "",
        f"**{sections['overall_heading']}**",
        "",
        sections["overall_assessment"],
        "",
        f"**{sections['grade_label']}: {sections['grade']}**",
        "",
        "## Scorecard",
    ]
    for key, payload in review_scorecard.get("dimensions", {}).items():
        lines.append(f"- {key}: {payload.get('score')} / {payload.get('summary', '')}")
    for heading_key, items_key in [
        ("critical_heading", "critical_items"),
        ("major_heading", "major_items"),
        ("minor_heading", "minor_items"),
        ("recurring_heading", "recurring_items"),
    ]:
        lines.extend(["", f"## {sections[heading_key]}"])
        lines.extend([f"- {item}" for item in sections[items_key]])
    lines.extend(["", f"## {sections['style_heading']}", sections["style_text"], "", f"## {sections['next_heading']}"])
    lines.extend([f"- {item}" for item in sections["next_items"]])
    with open(fallback, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")
    return fallback


def main():
    if len(sys.argv) != 6:
        print(json.dumps({"error": "Usage: generate-cover-memo.py <review_manifest_json> <issue_registry_json> <review_scorecard_json> <verification_audit_json> <output_docx>"}))
        sys.exit(1)

    manifest = load_json(sys.argv[1])
    issue_registry = load_json(sys.argv[2])
    review_scorecard = load_json(sys.argv[3])
    verification_audit = load_json(sys.argv[4])
    output_docx = sys.argv[5]
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)

    sections = build_sections(manifest, issue_registry, review_scorecard, verification_audit)

    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.shared import Inches, Mm, Pt
    except Exception:
        fallback = write_markdown_fallback(output_docx, sections, review_scorecard)
        print(json.dumps({"output_path": fallback, "fallback": True}, ensure_ascii=False))
        return

    doc = Document()
    section = doc.sections[0]
    if str(sections["language"]).lower().startswith("ko"):
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    title = doc.add_heading(sections["title"], level=0)
    title.runs[0].font.size = Pt(18)

    for label, value in sections["identification"]:
        para = doc.add_paragraph()
        para.add_run(f"{label}: ").bold = True
        para.add_run(str(value))

    doc.add_heading(sections["release_heading"], level=1)
    para = doc.add_paragraph()
    run = para.add_run(sections["release"])
    run.bold = True
    run.font.size = Pt(14)
    if sections["release_rationale"]:
        doc.add_paragraph(sections["release_rationale"])

    doc.add_heading(sections["overall_heading"], level=1)
    doc.add_paragraph(sections["overall_assessment"])
    para = doc.add_paragraph()
    para.add_run(f"{sections['grade_label']}: ").bold = True
    para.add_run(sections["grade"])

    doc.add_heading(sections["scorecard_heading"], level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "Score"
    hdr[2].text = "Key Findings"
    for key, payload in review_scorecard.get("dimensions", {}).items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = "N/A" if payload.get("skipped") else str(payload.get("score"))
        row[2].text = payload.get("summary", "")

    for heading_key, items_key in [
        ("critical_heading", "critical_items"),
        ("major_heading", "major_items"),
        ("minor_heading", "minor_items"),
        ("recurring_heading", "recurring_items"),
    ]:
        doc.add_heading(sections[heading_key], level=1)
        for item in sections[items_key]:
            doc.add_paragraph(item, style="List Number")

    doc.add_heading(sections["style_heading"], level=1)
    doc.add_paragraph(sections["style_text"])
    doc.add_heading(sections["next_heading"], level=1)
    for item in sections["next_items"]:
        doc.add_paragraph(item, style="List Number")

    doc.save(output_docx)
    print(json.dumps({"output_path": output_docx, "fallback": False}, ensure_ascii=False))


if __name__ == "__main__":
    main()
